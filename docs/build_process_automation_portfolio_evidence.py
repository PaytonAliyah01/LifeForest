from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "lifeforest-process-automation-evidence.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(25, 25, 25)
MUTED = RGBColor(90, 90, 90)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
BORDER = "C9D3DF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_inches):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    for index, width in enumerate(widths_inches):
        table.columns[index].width = Inches(width)
        for cell in table.columns[index].cells:
            cell.width = Inches(width)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_cell_text(cell, text, bold=False, color=None, size=9.1):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.10
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    paragraph.add_run(text)
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.style = "Normal"
    paragraph.add_run(text)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.167
    paragraph.add_run(text)
    return paragraph


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.167
    paragraph.add_run(text)
    return paragraph


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    set_table_borders(table, color="D8E0EA", size="4")
    mark_header_row(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
    title_para = cell.paragraphs[0]
    title_para.paragraph_format.space_after = Pt(3)
    title_run = title_para.add_run(title)
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(10.5)
    title_run.font.bold = True
    title_run.font.color.rgb = DARK_BLUE
    body_para = cell.add_paragraph()
    body_para.paragraph_format.space_after = Pt(0)
    body_para.paragraph_format.line_spacing = 1.10
    body_run = body_para.add_run(text)
    body_run.font.name = "Calibri"
    body_run.font.size = Pt(10)
    body_run.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    mark_header_row(table.rows[0])
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_shading(header_cells[i], LIGHT_BLUE)
        set_cell_margins(header_cells[i])
        set_cell_text(header_cells[i], header, bold=True, color=DARK_BLUE, size=9.2)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_text(cells[i], value, size=8.85)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    for list_style_name in ["List Bullet", "List Number"]:
        style = styles[list_style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("LifeForest - Process Automation Evidence")
    footer_run.font.name = "Calibri"
    footer_run.font.size = Pt(8.5)
    footer_run.font.color.rgb = MUTED


def build_docx():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    title_run = title.add_run("Software Development Process Automation")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = BLUE

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle.add_run("Portfolio evidence: how I applied automation to the LifeForest development process")
    subtitle_run.font.name = "Calibri"
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = MUTED

    add_callout(
        doc,
        "Learning outcome claim",
        "I applied software development process automation in LifeForest by using GitHub Actions for CI, automated backend and frontend quality checks, dependency and security scanning, Docker containerization, Docker Compose startup, build artifacts, and scripts that make local development repeatable.",
    )

    add_heading(doc, "What Process Automation Means", 1)
    add_body(
        doc,
        "Software development process automation means using tools, scripts, and workflows to reduce manual work in the development lifecycle. Instead of manually checking every change, the project automatically builds, tests, scans, and packages important parts of the application.",
    )
    add_body(
        doc,
        "For LifeForest this matters because the project has multiple moving parts: a Spring Boot backend, a React Native/Expo frontend, a PostgreSQL database, authentication, routines, tasks, focus sessions, reflections, analytics, achievements, and forest progress. Automation lowers the chance that a rushed manual step breaks one of these parts unnoticed.",
    )

    add_heading(doc, "Automation I Applied in LifeForest", 1)
    add_matrix(
        doc,
        ["Automation area", "LifeForest implementation", "Value for the project"],
        [
            ["Version control", "Project is managed with Git and GitHub workflow triggers on push and pull request to main.", "Changes are traceable and automation runs before work is merged or accepted."],
            ["Backend CI", "GitHub Actions sets up Java 21, runs Gradle assemble, uploads the backend JAR artifact, and then runs Gradle tests.", "Confirms the Spring Boot API still builds and that backend tests pass."],
            ["Frontend CI", "GitHub Actions sets up Node 22, runs npm ci, Expo lint, and TypeScript type checking.", "Catches frontend code quality and type errors before they reach users."],
            ["Security automation", "Backend OWASP Dependency-Check and frontend npm audit run automatically, with high-risk frontend audit failures blocked.", "Reduces the risk of shipping known vulnerable dependencies."],
            ["Containerization", "Backend and frontend Dockerfiles define reproducible runtime environments.", "Makes setup more consistent across machines and supports deployment packaging."],
            ["Local environment", "Docker Compose starts PostgreSQL and backend services with environment variables.", "Reduces manual setup and makes the development environment repeatable."],
        ],
        [1.45, 2.75, 2.3],
    )

    add_heading(doc, "Automated CI/CD Workflow", 1)
    add_body(
        doc,
        "The main automation workflow is defined in .github/workflows/ci.yml. It runs when code is pushed to main or when a pull request targets main. The workflow is split into backend and frontend jobs so each part of the application is validated with the right toolchain.",
    )
    for item in [
        "Backend build job: checks out the code, installs Java 21, caches Gradle dependencies, runs ./gradlew assemble, and uploads the generated JAR artifact.",
        "Backend test job: waits for the backend build, runs ./gradlew test, runs ./gradlew dependencyCheckAnalyze, and uploads the OWASP Dependency-Check report.",
        "Frontend build job: checks out the code, installs Node 22, runs npm ci, runs npm run lint, and runs npm run typecheck.",
        "Frontend test/security job: installs dependencies, runs npm run test --if-present, and runs npm audit --audit-level=high.",
        "Concurrency control: the workflow cancels older in-progress runs for the same branch, which avoids wasting time on outdated checks.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Before and After Process", 1)
    add_matrix(
        doc,
        ["Development step", "Without automation", "With LifeForest automation"],
        [
            ["Build validation", "Developer manually runs builds and may forget one side of the project.", "GitHub Actions builds the backend and validates frontend dependencies automatically."],
            ["Testing", "Tests depend on the developer remembering to run them locally.", "Backend tests run in CI after backend build succeeds."],
            ["Frontend quality", "Lint and TypeScript errors may only appear during manual review or runtime.", "CI runs Expo lint and TypeScript type checking."],
            ["Security checks", "Dependency vulnerabilities are checked manually or not checked regularly.", "OWASP Dependency-Check and npm audit are part of the automated process."],
            ["Environment setup", "Developers manually install and configure database/backend services.", "Docker Compose and startup scripts make setup repeatable."],
            ["Release preparation", "Build artifacts have to be created manually.", "The backend build artifact is generated and uploaded by CI."],
        ],
        [1.5, 2.45, 2.55],
    )

    add_heading(doc, "Containerization and Repeatable Setup", 1)
    add_body(
        doc,
        "LifeForest uses Docker to package application environments. The backend Dockerfile builds the Spring Boot application with Java 21 and then runs the produced JAR in a smaller Java runtime image. The frontend Dockerfile installs Node, dependencies, and starts Expo on a predictable port.",
    )
    add_body(
        doc,
        "The compose.yaml file defines PostgreSQL and backend services. It also centralizes database credentials and JWT configuration through environment variables. This means the project can be started in a predictable way instead of relying on hidden local machine settings.",
    )
    add_body(
        doc,
        "The start-backend-local.ps1 script supports local development by importing .env values, applying safe defaults, starting only the database container when needed, and running the backend with Gradle bootRun. This automates a repeated developer task and reduces setup mistakes.",
    )

    add_heading(doc, "Security and Quality Automation", 1)
    add_matrix(
        doc,
        ["Check", "Tool or command", "Purpose"],
        [
            ["Backend tests", "./gradlew test", "Validates backend service, controller, and integration behaviour."],
            ["Backend build", "./gradlew assemble / bootJar", "Confirms the Spring Boot API compiles and packages correctly."],
            ["Backend dependency scan", "./gradlew dependencyCheckAnalyze", "Produces HTML and JSON OWASP reports and fails for CVSS 7.0 or higher."],
            ["Frontend lint", "npm run lint", "Checks Expo/React Native code style and likely mistakes."],
            ["Frontend type check", "npm run typecheck", "Uses TypeScript to catch type errors before runtime."],
            ["Frontend dependency audit", "npm audit --audit-level=high", "Blocks high-severity dependency vulnerabilities."],
            ["Optional dynamic scan", "security-check.ps1 with ZapTarget", "Can run OWASP ZAP baseline scan against a running backend."],
        ],
        [1.55, 2.2, 2.75],
    )

    add_heading(doc, "How This Changes the Development Process", 1)
    for item in [
        "Errors are found earlier because builds, tests, linting, type checking, and audits run automatically.",
        "Manual work is reduced because common tasks are scripted instead of repeated from memory.",
        "Team collaboration becomes safer because pull requests and pushes receive consistent validation.",
        "Security is treated as part of the normal workflow instead of a separate last-minute activity.",
        "Deployment preparation becomes more reliable because Dockerfiles and CI artifacts describe how the application is built and run.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Evidence From My Repository", 1)
    add_body(
        doc,
        "The strongest evidence is the CI workflow in .github/workflows/ci.yml, the backend and frontend Dockerfiles, compose.yaml, security-check.ps1, start-backend-local.ps1, frontend package scripts, and the OWASP Dependency-Check configuration in backend/build.gradle.",
    )
    add_body(
        doc,
        "These files show that I did not only describe process automation theoretically. I applied it to LifeForest by automating builds, tests, linting, type checking, dependency scanning, artifact creation, container setup, and repeated local startup tasks.",
    )

    add_heading(doc, "Conclusion", 1)
    add_body(
        doc,
        "Software development process automation improved the LifeForest workflow by making quality checks repeatable and less dependent on manual discipline. The project can now detect build errors, failing tests, frontend quality issues, type errors, and vulnerable dependencies earlier in the process.",
    )
    add_body(
        doc,
        "The main improvement is reliability. Instead of trusting that every developer remembered every command, LifeForest uses automated workflows and scripts to enforce a consistent process. This supports faster development, safer collaboration, and a better chance of delivering a stable application to users.",
    )

    doc.core_properties.author = "LifeForest"
    doc.core_properties.title = "LifeForest Software Development Process Automation Evidence"
    doc.core_properties.subject = "Portfolio evidence for software development process automation"
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_docx()
    print(DOCX_PATH)
