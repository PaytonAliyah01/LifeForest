from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    KeepTogether,
    ListFlowable,
    ListItem,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "testing-portfolio-evidence.docx"
PDF_PATH = OUT_DIR / "testing-portfolio-evidence.pdf"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(25, 25, 25)
MUTED = RGBColor(90, 90, 90)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
BORDER = "C9D3DF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text, bold=False, color=None, size=9.4):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    paragraph.add_run(text)
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.style = "Normal"
    paragraph.add_run(text)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.add_run(text)
    return paragraph


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.25)
    set_table_borders(table, color="D8E0EA", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
    title_para = cell.paragraphs[0]
    title_para.paragraph_format.space_after = Pt(3)
    title_run = title_para.add_run(title)
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(10.5)
    title_run.font.bold = True
    title_run.font.color.rgb = DARK_BLUE
    body_para = cell.add_paragraph()
    body_para.paragraph_format.space_after = Pt(0)
    body_run = body_para.add_run(text)
    body_run.font.name = "Calibri"
    body_run.font.size = Pt(10)
    body_run.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.7)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 15.5, BLUE, 14, 7),
        ("Heading 2", 12.5, BLUE, 10, 5),
        ("Heading 3", 11.5, DARK_BLUE, 8, 3),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(2)
    title_run = title.add_run("Testing")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = BLUE

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle.add_run("Portfolio evidence: how I applied the testing workshop to LifeForest")
    subtitle_run.font.name = "Calibri"
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = MUTED

    add_callout(
        doc,
        "Learning outcome claim",
        "I applied testing in LifeForest by choosing high-value and high-risk areas to test, writing automated backend unit and integration tests, using mocks and an H2 test database, checking frontend quality with linting and TypeScript, and automating the process in GitHub Actions.",
    )

    add_heading(doc, "What Testing Means", 1)
    add_body(
        doc,
        "Testing is the process of validating that software works as intended and that important changes do not break existing behaviour. The workshop explains that testing is about quality and trust: it reduces bugs, improves confidence, and proves that the application is working as ordered.",
    )
    add_body(
        doc,
        "For LifeForest, testing matters because the application combines a mobile frontend, Spring Boot API, PostgreSQL data, authentication, routines, tasks, focus sessions, reflections, analytics, achievements, and forest progress. If one part breaks, the user experience and collected productivity data can become unreliable.",
    )

    add_heading(doc, "My Testing Strategy", 1)
    for item in [
        "I focused automated tests on backend logic and API behaviour because these areas protect the main project value: reliable routines, tasks, focus sessions, analytics, achievements, reflections, users, and trees.",
        "I used the test pyramid by placing many checks at the unit/service level and adding integration/controller tests where the API, Spring context, repositories, and database behaviour needed to work together.",
        "I used mocks for service tests so dependencies such as repositories could be controlled and edge cases could be tested quickly.",
        "I used H2 with create-drop schema settings for tests, so database-backed integration tests can run without depending on the development PostgreSQL database.",
        "I added CI checks so tests and quality gates run automatically after code changes instead of relying only on manual testing.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Evidence From LifeForest", 1)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [1.35, 2.15, 1.7, 1.35]
    for index, width in enumerate(widths):
        table.columns[index].width = Inches(width)
    set_table_borders(table)
    headers = ["Testing idea", "How I applied it", "Evidence location", "Why it matters"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        set_cell_text(cell, header, bold=True, color=DARK_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    rows = [
        (
            "Unit tests",
            "Service tests validate business logic with JUnit and Mockito, for example task linking/deleting and analytics calculations.",
            "backend/src/test/java",
            "Fast feedback on core logic.",
        ),
        (
            "Mocking",
            "Repository dependencies are mocked in service tests so success, empty data, and error paths can be controlled.",
            "TaskServiceTest, AnalyticsServiceTest",
            "Tests stay isolated and repeatable.",
        ),
        (
            "Integration tests",
            "Controller integration tests use Spring Boot, MockMvc, repositories, transactions, and H2.",
            "TreeControllerIntegrationTest and others",
            "Checks API and persistence behaviour together.",
        ),
        (
            "Regression testing",
            "There are 17 backend test files with 52 test methods that can be rerun after changes.",
            "backend/src/test/java",
            "Prevents old behaviour from breaking unnoticed.",
        ),
        (
            "Frontend checks",
            "The Expo frontend has linting and TypeScript type checking commands.",
            "frontend/package.json",
            "Catches UI code mistakes before runtime.",
        ),
        (
            "Automated process",
            "GitHub Actions builds the backend, runs backend tests, runs OWASP Dependency-Check, lints/typechecks frontend, and runs npm audit.",
            ".github/workflows/ci.yml",
            "Quality checks run consistently in the pipeline.",
        ),
        (
            "Manual testing",
            "I manually checked important app flows such as login, tasks, routines, focus sessions, reflections, achievements, analytics, and forest progress.",
            "App behaviour and screenshots/notes",
            "Confirms the experience from a user perspective.",
        ),
    ]
    for row_data in rows:
        row = table.add_row()
        for index, text in enumerate(row_data):
            cell = row.cells[index]
            set_cell_margins(cell)
            set_cell_text(cell, text, bold=(index == 0), color=INK)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    add_heading(doc, "Examples Of Test Quality", 1)
    add_body(
        doc,
        "TaskServiceTest checks that a task is trimmed, saved, linked to the correct routine, and removed from the routine before deletion. This is better than only checking that the method returns something, because it validates side effects and relationships.",
    )
    add_body(
        doc,
        "AnalyticsServiceTest checks normal calculations, empty data, and an unknown user. This covers valid input, no-data boundaries, and an error condition, which connects to the workshop idea of choosing meaningful cases instead of only testing the happy path.",
    )
    add_body(
        doc,
        "TreeControllerIntegrationTest creates two users and verifies that the API only returns the selected user's forest. This is valuable because data separation is a high-risk area for a multi-user application.",
    )

    add_heading(doc, "Reflection", 1)
    add_body(
        doc,
        "The workshop changed my approach from 'I ran it and it worked' to a more reliable process. I now use automated tests for important backend behaviour, static checks for frontend code, and CI automation so checks happen repeatedly.",
    )
    add_body(
        doc,
        "The strongest part of my testing evidence is the backend coverage for services and controllers. The main improvement I still need is more automated frontend testing, especially for user flows such as logging in, creating a routine, starting a focus session, submitting a reflection, and viewing updated analytics or forest progress.",
    )

    add_heading(doc, "Next Improvements", 1)
    for item in [
        "Add frontend unit tests with Jest or Vitest for components and hooks that contain logic.",
        "Add end-to-end tests with Playwright, Cypress, or Detox for the most important mobile/web flows.",
        "Add more security-focused tests for authentication, authorization, invalid JWTs, and user ownership rules.",
        "Keep screenshots or CI run links in Portflow as extra proof next to this document.",
    ]:
        add_bullet(doc, item)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("LifeForest - Testing Portfolio Evidence")
    footer_run.font.name = "Calibri"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = MUTED

    doc.save(DOCX_PATH)
    return DOCX_PATH


def make_pdf_styles():
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle("Title", parent=base["Title"], fontName="Helvetica-Bold", fontSize=24, leading=29, textColor=colors.HexColor("#2E74B5"), alignment=TA_LEFT, spaceAfter=2),
        "subtitle": ParagraphStyle("Subtitle", parent=base["Normal"], fontName="Helvetica", fontSize=10.7, leading=14, textColor=colors.HexColor("#5A5A5A"), spaceAfter=12),
        "h1": ParagraphStyle("Heading1", parent=base["Heading1"], fontName="Helvetica-Bold", fontSize=15.5, leading=19, textColor=colors.HexColor("#2E74B5"), spaceBefore=13, spaceAfter=7),
        "body": ParagraphStyle("Body", parent=base["BodyText"], fontName="Helvetica", fontSize=10.25, leading=13.2, textColor=colors.HexColor("#191919"), spaceAfter=6),
        "small": ParagraphStyle("Small", parent=base["BodyText"], fontName="Helvetica", fontSize=8.2, leading=10.2, textColor=colors.HexColor("#191919"), spaceAfter=0),
        "small_bold": ParagraphStyle("SmallBold", parent=base["BodyText"], fontName="Helvetica-Bold", fontSize=8.2, leading=10.2, textColor=colors.HexColor("#191919"), spaceAfter=0),
        "header": ParagraphStyle("Header", parent=base["BodyText"], fontName="Helvetica-Bold", fontSize=8.6, leading=10.3, textColor=colors.HexColor("#1F4D78"), spaceAfter=0),
        "callout_title": ParagraphStyle("CalloutTitle", parent=base["BodyText"], fontName="Helvetica-Bold", fontSize=10.2, leading=12.5, textColor=colors.HexColor("#1F4D78"), spaceAfter=3),
        "callout_body": ParagraphStyle("CalloutBody", parent=base["BodyText"], fontName="Helvetica", fontSize=9.7, leading=12.3, textColor=colors.HexColor("#191919"), spaceAfter=0),
    }


def pdf_bullets(items, styles):
    return ListFlowable(
        [ListItem(Paragraph(item, styles["body"]), leftIndent=8) for item in items],
        bulletType="bullet",
        start="circle",
        leftIndent=17,
        bulletFontName="Helvetica",
        bulletFontSize=7,
        bulletOffsetY=1,
        spaceAfter=4,
    )


def pdf_callout(title, text, styles):
    table = Table(
        [[Paragraph(title, styles["callout_title"]), Paragraph(text, styles["callout_body"])]],
        colWidths=[1.45 * inch, 4.75 * inch],
        hAlign="LEFT",
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F4F6F9")),
                ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#C9D3DF")),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )
    return KeepTogether([table, Spacer(1, 8)])


def pdf_evidence_table(styles):
    headers = ["Testing idea", "How I applied it", "Evidence location", "Why it matters"]
    rows = [
        ["Unit tests", "Service tests validate business logic with JUnit and Mockito.", "backend/src/test/java", "Fast feedback on core logic."],
        ["Mocking", "Repository dependencies are mocked to control success, empty data, and error paths.", "TaskServiceTest, AnalyticsServiceTest", "Isolated and repeatable tests."],
        ["Integration tests", "Controller integration tests use Spring Boot, MockMvc, repositories, transactions, and H2.", "TreeControllerIntegrationTest and others", "API and persistence checked together."],
        ["Regression testing", "17 backend test files with 52 test methods can be rerun after changes.", "backend/src/test/java", "Old behaviour is protected."],
        ["Frontend checks", "Expo linting and TypeScript type checking are available.", "frontend/package.json", "Catches UI mistakes early."],
        ["Automated process", "GitHub Actions builds, tests, checks dependencies, lints, typechecks, and audits.", ".github/workflows/ci.yml", "Checks run consistently."],
        ["Manual testing", "Login, tasks, routines, focus sessions, reflections, achievements, analytics, and forest progress were checked manually.", "App behaviour and notes", "Confirms user-facing flow."],
    ]
    data = [[Paragraph(h, styles["header"]) for h in headers]]
    for row in rows:
        data.append(
            [
                Paragraph(row[0], styles["small_bold"]),
                Paragraph(row[1], styles["small"]),
                Paragraph(row[2], styles["small"]),
                Paragraph(row[3], styles["small"]),
            ]
        )
    table = Table(data, colWidths=[1.15 * inch, 2.08 * inch, 1.55 * inch, 1.42 * inch], repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
                ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#C9D3DF")),
                ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#C9D3DF")),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )
    return table


def pdf_footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.HexColor("#5A5A5A"))
    canvas.drawRightString(7.6 * inch, 0.45 * inch, "LifeForest - Testing Portfolio Evidence")
    canvas.restoreState()


def build_pdf():
    styles = make_pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        leftMargin=0.9 * inch,
        rightMargin=0.9 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.72 * inch,
        title="Testing Portfolio Evidence",
        author="LifeForest",
    )

    story = [
        Paragraph("Testing", styles["title"]),
        Paragraph("Portfolio evidence: how I applied the testing workshop to LifeForest", styles["subtitle"]),
        pdf_callout(
            "Learning outcome claim",
            "I applied testing in LifeForest by choosing high-value and high-risk areas to test, writing automated backend unit and integration tests, using mocks and an H2 test database, checking frontend quality with linting and TypeScript, and automating the process in GitHub Actions.",
            styles,
        ),
        Paragraph("What Testing Means", styles["h1"]),
        Paragraph("Testing validates that software works as intended and that important changes do not break existing behaviour. It is about quality and trust: fewer bugs, more confidence, and stronger proof that the project works.", styles["body"]),
        Paragraph("For LifeForest, this matters because the project combines a mobile frontend, Spring Boot API, database behaviour, authentication, routines, tasks, focus sessions, reflections, analytics, achievements, and forest progress.", styles["body"]),
        Paragraph("My Testing Strategy", styles["h1"]),
        pdf_bullets(
            [
                "Focused automated tests on high-value backend logic and API behaviour.",
                "Used the test pyramid: many unit/service tests plus integration/controller tests for combined behaviour.",
                "Used mocks to isolate service logic and test edge cases quickly.",
                "Used H2 with create-drop settings for database-backed integration tests.",
                "Used CI so tests and quality checks run automatically after changes.",
            ],
            styles,
        ),
        Paragraph("Evidence From LifeForest", styles["h1"]),
        pdf_evidence_table(styles),
        Paragraph("Examples Of Test Quality", styles["h1"]),
        Paragraph("TaskServiceTest checks task trimming, saving, routine linking, and delete side effects. AnalyticsServiceTest checks normal calculations, empty data, and unknown-user errors. TreeControllerIntegrationTest verifies user-specific forest data through the API and database layer.", styles["body"]),
        Paragraph("Reflection", styles["h1"]),
        Paragraph("The workshop changed my approach from 'I ran it and it worked' to a repeatable process. I now use automated backend tests, frontend static checks, and CI automation to catch problems earlier.", styles["body"]),
        Paragraph("The strongest current evidence is backend service and controller coverage. The main next step is adding automated frontend tests and end-to-end tests for the most important user flows.", styles["body"]),
        Paragraph("Next Improvements", styles["h1"]),
        pdf_bullets(
            [
                "Add frontend unit tests with Jest or Vitest for components and hooks.",
                "Add end-to-end tests with Playwright, Cypress, or Detox for important flows.",
                "Add more security-focused tests for authentication, authorization, invalid JWTs, and user ownership rules.",
                "Keep screenshots or CI run links in Portflow as extra proof next to this document.",
            ],
            styles,
        ),
    ]

    doc.build(story, onFirstPage=pdf_footer, onLaterPages=pdf_footer)
    return PDF_PATH


if __name__ == "__main__":
    print(build_docx())
    print(build_pdf())
