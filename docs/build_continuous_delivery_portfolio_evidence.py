from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "continuous-delivery-portfolio-evidence.docx"
PDF_PATH = OUT_DIR / "continuous-delivery-portfolio-evidence.pdf"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(25, 25, 25)
MUTED = RGBColor(90, 90, 90)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
BORDER = "C9D3DF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=130, bottom=90, end=130):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_inches):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    for index, width in enumerate(widths_inches):
        table.columns[index].width = Inches(width)
        for cell in table.columns[index].cells:
            cell.width = Inches(width)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_cell_text(cell, text, bold=False, color=None, size=8.9):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    paragraph.add_run(text)


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.style = "Normal"
    paragraph.add_run(text)


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.167
    paragraph.add_run(text)


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    set_table_borders(table, color="D8E0EA", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
    title_para = cell.paragraphs[0]
    title_para.paragraph_format.space_after = Pt(3)
    title_run = title_para.add_run(title)
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(10.5)
    title_run.font.bold = True
    title_run.font.color.rgb = DARK_BLUE
    body_para = cell.add_paragraph()
    body_para.paragraph_format.space_after = Pt(0)
    body_run = body_para.add_run(text)
    body_run.font.name = "Calibri"
    body_run.font.size = Pt(10)
    body_run.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    mark_header_row(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        set_cell_text(cell, header, bold=True, color=DARK_BLUE, size=9.1)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row_data in rows:
        row = table.add_row()
        for index, value in enumerate(row_data):
            cell = row.cells[index]
            set_cell_margins(cell)
            set_cell_text(cell, value, bold=(index == 0), color=INK)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("LifeForest - Continuous Delivery Evidence")
    footer_run.font.name = "Calibri"
    footer_run.font.size = Pt(8.5)
    footer_run.font.color.rgb = MUTED


def build_docx():
    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    title_run = title.add_run("Continuous Delivery")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = BLUE

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle.add_run("Portfolio evidence: how I applied CD principles to LifeForest")
    subtitle_run.font.name = "Calibri"
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = MUTED

    add_callout(
        doc,
        "Learning outcome claim",
        "I applied Continuous Delivery in LifeForest by keeping the project in a releasable state through automated build, test, lint, type-check, dependency-audit, artifact, Docker, and deployment-preparation steps. Production deployment is still a manual decision, so the project demonstrates Continuous Delivery rather than full Continuous Deployment.",
    )

    add_heading(doc, "What Continuous Delivery Means", 1)
    add_body(
        doc,
        "Continuous Delivery means that code changes are automatically validated and prepared for release. The release to production is still controlled manually, but the project should always be close to a deployable state.",
    )
    add_body(
        doc,
        "For LifeForest this is important because the project has a Spring Boot backend, Expo/React Native frontend, PostgreSQL database, Docker setup, security checks, and deployment instructions. CD reduces release stress by checking changes earlier and making deployment preparation repeatable.",
    )

    add_heading(doc, "CD Pipeline Evidence", 1)
    add_matrix(
        doc,
        ["CD practice", "LifeForest implementation", "Evidence location", "Release value"],
        [
            ["Automated trigger", "Pipeline runs on push to main and pull requests targeting main.", ".github/workflows/ci.yml", "Every important change is validated automatically."],
            ["Backend build", "GitHub Actions sets up Java 21, caches Gradle, runs ./gradlew assemble, and uploads the JAR.", ".github/workflows/ci.yml", "Backend is packaged as a releasable artifact."],
            ["Backend tests", "Backend test job runs ./gradlew test after build succeeds.", "backend/src/test and CI", "Business/API behaviour is checked before release."],
            ["Dependency security", "Gradle OWASP Dependency-Check runs in CI and uploads reports.", "backend/build.gradle and CI", "Known risky dependencies are visible before release."],
            ["Frontend validation", "CI runs npm ci, npm run lint, npm run typecheck, optional tests, and npm audit.", "frontend/package.json and CI", "Frontend mistakes are caught before packaging/release."],
            ["Docker packaging", "Backend and frontend Dockerfiles define repeatable build/runtime environments.", "backend/Dockerfile and frontend/Dockerfile", "Builds are more portable and predictable."],
            ["Manual release control", "README explains Docker Hub push and VM deployment with docker compose pull/up.", "README.md", "Release remains deliberate, matching CD."],
        ],
        [1.28, 2.25, 1.5, 1.47],
    )

    add_heading(doc, "Continuous Delivery, Not Continuous Deployment", 1)
    add_matrix(
        doc,
        ["Aspect", "What LifeForest does", "What this shows"],
        [
            ["Automatic checks", "Build, tests, linting, type checks, dependency checks, and audits are automated.", "The project is prepared for release continuously."],
            ["Production release", "Docker image push and VM deployment are documented but manually triggered.", "The team keeps control over when the release goes live."],
            ["Risk management", "CI fails when important checks fail, while deployment is still a conscious decision.", "This matches Continuous Delivery rather than Continuous Deployment."],
        ],
        [1.4, 2.95, 2.15],
    )

    add_heading(doc, "Docker and Repeatable Release Preparation", 1)
    for item in [
        "The backend Dockerfile builds the Spring Boot JAR with Java 21 and then runs it in a smaller Java runtime image.",
        "The frontend Dockerfile installs Node, dependencies, and starts Expo on predictable ports.",
        "compose.yaml starts PostgreSQL and the backend with environment variables for database and JWT configuration.",
        "The README documents Docker-based startup, Docker Hub image tagging/pushing, and VM deployment with docker compose pull and docker compose up -d.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Release Flow I Can Demonstrate", 1)
    add_matrix(
        doc,
        ["Step", "LifeForest command or file", "Purpose"],
        [
            ["1. Change code", "Git push or pull request to main", "Starts automated validation."],
            ["2. Validate backend", "./gradlew assemble and ./gradlew test in CI", "Builds and tests the Spring Boot API."],
            ["3. Validate frontend", "npm ci, lint, typecheck, test if present, npm audit", "Checks frontend quality and dependencies."],
            ["4. Prepare artifact", "backend-build-artifacts uploaded by GitHub Actions", "Creates a backend artifact that can be released."],
            ["5. Package environment", "Dockerfiles and compose.yaml", "Makes runtime setup reproducible."],
            ["6. Manual release", "docker build, docker push, docker compose pull/up", "Promotes a tested build when the team decides."],
        ],
        [0.95, 2.45, 3.1],
    )

    add_heading(doc, "Reflection", 1)
    add_body(
        doc,
        "This metric helped me understand that deployment readiness is not only about writing code. A project also needs automated checks, repeatable packaging, clear environment configuration, and a release path that can be followed without guessing.",
    )
    add_body(
        doc,
        "The strongest CD evidence in LifeForest is the GitHub Actions workflow combined with Dockerfiles, Docker Compose, and documented VM deployment commands. The biggest next improvement would be automatically building and publishing Docker images after successful CI, then adding a staging environment before manual production release.",
    )

    add_heading(doc, "Next Improvements", 1)
    for item in [
        "Add a GitHub Actions job that builds Docker images after tests pass.",
        "Push versioned images to Docker Hub or GitHub Container Registry automatically after approval.",
        "Add a staging environment that runs the same image intended for production.",
        "Add release tags and changelog notes so each production candidate is traceable.",
        "Add health checks and smoke tests after starting the Docker Compose environment.",
    ]:
        add_bullet(doc, item)

    doc.core_properties.author = "LifeForest"
    doc.core_properties.title = "LifeForest Continuous Delivery Evidence"
    doc.core_properties.subject = "Portfolio evidence for Continuous Delivery"
    doc.save(DOCX_PATH)
    return DOCX_PATH


def styles():
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle("Title", parent=base["Title"], fontName="Helvetica-Bold", fontSize=24, leading=29, textColor=colors.HexColor("#2E74B5"), alignment=TA_LEFT, spaceAfter=2),
        "subtitle": ParagraphStyle("Subtitle", parent=base["Normal"], fontName="Helvetica", fontSize=10.7, leading=14, textColor=colors.HexColor("#5A5A5A"), spaceAfter=12),
        "h1": ParagraphStyle("Heading1", parent=base["Heading1"], fontName="Helvetica-Bold", fontSize=15, leading=18, textColor=colors.HexColor("#2E74B5"), spaceBefore=12, spaceAfter=6),
        "body": ParagraphStyle("Body", parent=base["BodyText"], fontName="Helvetica", fontSize=9.8, leading=12.6, textColor=colors.HexColor("#191919"), spaceAfter=5),
        "small": ParagraphStyle("Small", parent=base["BodyText"], fontName="Helvetica", fontSize=7.6, leading=9.4, textColor=colors.HexColor("#191919"), spaceAfter=0),
        "small_bold": ParagraphStyle("SmallBold", parent=base["BodyText"], fontName="Helvetica-Bold", fontSize=7.6, leading=9.4, textColor=colors.HexColor("#191919"), spaceAfter=0),
        "header": ParagraphStyle("Header", parent=base["BodyText"], fontName="Helvetica-Bold", fontSize=7.8, leading=9.5, textColor=colors.HexColor("#1F4D78"), spaceAfter=0),
        "callout_title": ParagraphStyle("CalloutTitle", parent=base["BodyText"], fontName="Helvetica-Bold", fontSize=9.8, leading=12, textColor=colors.HexColor("#1F4D78"), spaceAfter=2),
        "callout_body": ParagraphStyle("CalloutBody", parent=base["BodyText"], fontName="Helvetica", fontSize=9.2, leading=11.7, textColor=colors.HexColor("#191919"), spaceAfter=0),
    }


def pdf_bullets(items, sty):
    return ListFlowable(
        [ListItem(Paragraph(item, sty["body"]), leftIndent=8) for item in items],
        bulletType="bullet",
        leftIndent=16,
        bulletFontName="Helvetica",
        bulletFontSize=7,
        bulletOffsetY=1,
        spaceAfter=2,
    )


def pdf_table(headers, rows, widths, sty):
    data = [[Paragraph(h, sty["header"]) for h in headers]]
    for row in rows:
        data.append([Paragraph(row[0], sty["small_bold"])] + [Paragraph(value, sty["small"]) for value in row[1:]])
    table = Table(data, colWidths=[width * inch for width in widths], repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
        ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#C9D3DF")),
        ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#C9D3DF")),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    return table


def pdf_callout(title, text, sty):
    table = Table([[Paragraph(title, sty["callout_title"]), Paragraph(text, sty["callout_body"])]], colWidths=[1.35 * inch, 4.85 * inch])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F4F6F9")),
        ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#C9D3DF")),
        ("LEFTPADDING", (0, 0), (-1, -1), 9),
        ("RIGHTPADDING", (0, 0), (-1, -1), 9),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    return table


def footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.HexColor("#5A5A5A"))
    canvas.drawRightString(7.6 * inch, 0.45 * inch, "LifeForest - Continuous Delivery Evidence")
    canvas.restoreState()


def build_pdf():
    sty = styles()
    doc = SimpleDocTemplate(str(PDF_PATH), pagesize=letter, leftMargin=0.9 * inch, rightMargin=0.9 * inch, topMargin=0.72 * inch, bottomMargin=0.72 * inch)
    story = [
        Paragraph("Continuous Delivery", sty["title"]),
        Paragraph("Portfolio evidence: how I applied CD principles to LifeForest", sty["subtitle"]),
        pdf_callout("Learning outcome claim", "I applied Continuous Delivery in LifeForest by keeping the project in a releasable state through automated build, test, lint, type-check, dependency-audit, artifact, Docker, and deployment-preparation steps. Production deployment is still manual, so this is CD rather than full Continuous Deployment.", sty),
        Spacer(1, 8),
        Paragraph("What Continuous Delivery Means", sty["h1"]),
        Paragraph("Continuous Delivery means code changes are automatically validated and prepared for release. The final production release remains a manual decision, but the project stays close to deployable.", sty["body"]),
        Paragraph("CD Pipeline Evidence", sty["h1"]),
        pdf_table(
            ["CD practice", "LifeForest implementation", "Evidence location", "Release value"],
            [
                ["Automated trigger", "Runs on push to main and pull requests targeting main.", ".github/workflows/ci.yml", "Important changes are validated."],
                ["Backend build", "Sets up Java 21, runs Gradle assemble, uploads JAR.", ".github/workflows/ci.yml", "Backend becomes a releasable artifact."],
                ["Backend tests", "Runs ./gradlew test after build succeeds.", "backend/src/test and CI", "Business/API behaviour is checked."],
                ["Dependency security", "Runs OWASP Dependency-Check and uploads reports.", "backend/build.gradle and CI", "Risky dependencies are visible."],
                ["Frontend validation", "Runs npm ci, lint, typecheck, optional tests, npm audit.", "frontend/package.json and CI", "Frontend quality is checked."],
                ["Docker packaging", "Dockerfiles define repeatable environments.", "backend/Dockerfile and frontend/Dockerfile", "Builds are portable."],
                ["Manual release control", "Docker Hub push and VM deployment are documented.", "README.md", "Release timing stays controlled."],
            ],
            [1.1, 2.0, 1.45, 1.65],
            sty,
        ),
        Paragraph("CD, Not Continuous Deployment", sty["h1"]),
        pdf_table(
            ["Aspect", "What LifeForest does", "What this shows"],
            [
                ["Automatic checks", "Build, tests, linting, type checks, dependency checks, and audits are automated.", "The project is prepared for release continuously."],
                ["Production release", "Docker image push and VM deployment are documented but manually triggered.", "The team controls when the release goes live."],
                ["Risk management", "CI fails on important checks while deployment stays deliberate.", "This matches Continuous Delivery."],
            ],
            [1.25, 2.7, 2.25],
            sty,
        ),
        Paragraph("Docker and Release Preparation", sty["h1"]),
        pdf_bullets([
            "Backend Dockerfile builds the Spring Boot JAR with Java 21 and runs it in a smaller runtime image.",
            "Frontend Dockerfile installs Node dependencies and starts Expo on predictable ports.",
            "compose.yaml starts PostgreSQL and backend with environment-based configuration.",
            "README documents Docker startup, Docker Hub image pushing, and VM deployment.",
        ], sty),
        Paragraph("Reflection and Next Improvements", sty["h1"]),
        Paragraph("The strongest CD evidence is the GitHub Actions workflow combined with Dockerfiles, Docker Compose, and documented VM deployment commands. The main improvement would be automatically building and publishing Docker images after successful CI, then adding a staging environment before manual production release.", sty["body"]),
        pdf_bullets([
            "Add a GitHub Actions job that builds Docker images after tests pass.",
            "Push versioned images to Docker Hub or GitHub Container Registry after approval.",
            "Add a staging environment and smoke tests against the same image intended for production.",
        ], sty),
    ]
    doc.build(story, onFirstPage=footer, onLaterPages=footer)
    return PDF_PATH


if __name__ == "__main__":
    print(build_docx())
    print(build_pdf())
