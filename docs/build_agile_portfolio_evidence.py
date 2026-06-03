from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "lifeforest-agile-evidence.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(25, 25, 25)
MUTED = RGBColor(90, 90, 90)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
BORDER = "C9D3DF"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=130, bottom=90, end=130):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    for index, width in enumerate(widths):
        table.columns[index].width = Inches(width)
    tbl_grid = table._tbl.tblGrid
    for index, width in enumerate(widths):
        grid_col = tbl_grid.gridCol_lst[index]
        grid_col.set(qn("w:w"), str(int(width * 1440)))
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(widths[index] * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, color=None, size=9.5):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.10
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color or INK


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    paragraph.add_run(text)
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.style = "Normal"
    paragraph.add_run(text)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.add_run(text)
    return paragraph


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [6.5])
    set_table_borders(table, color="D8E0EA", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, top=155, bottom=155, start=180, end=180)
    title_para = cell.paragraphs[0]
    title_para.paragraph_format.space_after = Pt(3)
    title_run = title_para.add_run(title)
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(10.5)
    title_run.font.bold = True
    title_run.font.color.rgb = DARK_BLUE
    body_para = cell.add_paragraph()
    body_para.paragraph_format.space_after = Pt(0)
    body_run = body_para.add_run(text)
    body_run.font.name = "Calibri"
    body_run.font.size = Pt(10)
    body_run.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths, first_col_bold=True):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, widths)
    set_table_borders(table)
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        set_cell_text(cell, header, bold=True, color=DARK_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row_data in rows:
        row = table.add_row()
        for index, text in enumerate(row_data):
            cell = row.cells[index]
            set_cell_shading(cell, WHITE)
            set_cell_margins(cell)
            set_cell_text(cell, text, bold=(first_col_bold and index == 0), color=INK)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def configure_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    bullet = styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet.font.size = Pt(11)
    bullet.paragraph_format.left_indent = Inches(0.5)
    bullet.paragraph_format.first_line_indent = Inches(-0.25)
    bullet.paragraph_format.space_after = Pt(4)
    bullet.paragraph_format.line_spacing = 1.15

    return doc


def add_footer(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("LifeForest - Agile Portfolio Evidence")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED


def build_docx():
    doc = configure_document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(2)
    title_run = title.add_run("Agile Software Development")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = BLUE

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle.add_run(
        "Portfolio evidence: how I applied Agile working methods in the LifeForest individual project"
    )
    subtitle_run.font.name = "Calibri"
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = MUTED

    add_callout(
        doc,
        "Learning outcome claim",
        "I applied Agile by working iteratively, maintaining a visible backlog and progress log, prioritising user value, adapting the scope after technical feedback, using CI as a feedback loop, and delivering LifeForest in small increments instead of one large final build.",
    )

    add_heading(doc, "Project Context", 1)
    add_body(
        doc,
        "LifeForest is a mobile-first productivity and focus platform. The project contains an Expo React Native frontend, a Spring Boot backend API, and a PostgreSQL database. The product grew from a basic user foundation into authentication, routines, tasks, focus sessions, reflections, analytics, achievements, forest progress, security checks, Docker setup, and deployment documentation."
    )
    add_body(
        doc,
        "Because this was an individual project, I did not run a full Scrum team with a separate Scrum Master and Product Owner. Instead, I used a lightweight Scrum/Kanban-inspired approach: a backlog of work, short iterations, regular reflection, continuous reprioritisation, and a working product at the end of each increment."
    )

    add_heading(doc, "My Agile Method", 1)
    add_body(
        doc,
        "I chose a hybrid Agile method because the project needed flexibility. Scrum gave structure through goals, increments, and review/reflection moments. Kanban helped me keep work visible and move tasks through todo, in progress, testing, and done without forcing a full team ritual structure onto an individual project."
    )
    add_table(
        doc,
        ["Agile element", "How I applied it in LifeForest", "Evidence"],
        [
            (
                "Product goal",
                "Build a mobile-first platform that helps users plan routines, complete tasks, run focus sessions, reflect, and see progress through a forest metaphor.",
                "README.md and implemented frontend/backend features",
            ),
            (
                "Backlog",
                "I broke the work into feature areas: user foundation, authentication, frontend flows, CI, security, Docker/deployment, and later focus/routine/productivity features.",
                "LEARNING_TRACKER.txt phase and checklist views",
            ),
            (
                "Iterations",
                "I delivered in increments instead of waiting until the end. Each phase added a usable capability and left the project in a runnable state.",
                "Git history from March to May 2026",
            ),
            (
                "Review",
                "After each milestone I checked whether the feature worked, whether tests passed, and whether documentation/setup still matched the project.",
                "Postman checks, CI workflow, README updates",
            ),
            (
                "Retrospective",
                "I recorded issues and lessons learned so later work could avoid the same problems.",
                "LEARNING_TRACKER.txt troubleshooting and lessons sections",
            ),
            (
                "Continuous feedback",
                "The CI pipeline runs backend build/tests, dependency checks, frontend linting, type checking, and npm audit.",
                ".github/workflows/ci.yml",
            ),
        ],
        [1.35, 3.45, 1.70],
    )

    add_heading(doc, "Iterative Delivery Timeline", 1)
    add_body(
        doc,
        "The timeline below shows how the project evolved through small increments. This demonstrates the Agile principle of delivering working software frequently and improving the product based on what was learned in earlier increments."
    )
    add_table(
        doc,
        ["Increment", "Delivered value", "Adaptation or learning"],
        [
            (
                "Backend foundation",
                "User entity, repository, DTOs, mapper, service, controller, validation, persistence, exceptions.",
                "Created a clean layered architecture so later features could reuse the same structure.",
            ),
            (
                "Testing and API verification",
                "Unit tests, controller tests, edge-case tests, and manual Postman checks.",
                "Discovered and fixed endpoint, datasource, and validation issues before extending the app.",
            ),
            (
                "Authentication",
                "JWT login endpoint, invalid-credentials handling, frontend login/register integration, token storage.",
                "Moved from isolated user management to a real authenticated user flow.",
            ),
            (
                "Frontend UX and routing",
                "Welcome screen, login/register navigation, responsive layouts, native animation fix.",
                "Improved the product from API-only work toward a usable mobile experience.",
            ),
            (
                "Tooling, CI, and quality",
                "Expo upgrade, TypeScript fixes, backend/frontend CI jobs, linting, type checking, stable npm ci.",
                "Used automation to make feedback faster and reduce repeated manual checks.",
            ),
            (
                "Security and hardening",
                "Removed hardcoded secrets, environment placeholders, typed JWT config, Sonar/test cleanups.",
                "Adapted after quality/security feedback and reduced project risk.",
            ),
            (
                "Docker and deployment setup",
                "Backend/frontend Dockerfiles, Docker Compose, env-driven frontend API URL, beginner-friendly README.",
                "Made the project easier to run, test, and deploy in different environments.",
            ),
            (
                "Product feature expansion",
                "Routines, tasks, focus sessions, reflections, analytics, achievements, forest and tree progress.",
                "Expanded the product around the core user value: routines, focus, reflection, and visible progress.",
            ),
        ],
        [1.55, 2.70, 2.25],
    )

    add_heading(doc, "How I Used Agile Values", 1)
    for item in [
        "Individuals and interactions over processes and tools: I kept a simple progress tracker and made decisions based on what helped development move forward, instead of spending time on heavy process administration.",
        "Working software over comprehensive documentation: I used documentation to support the project, but prioritised runnable backend, frontend, database, CI, and Docker workflows.",
        "Customer collaboration over contract negotiation: I translated the user need into product value: routines, tasks, focus sessions, reflection, analytics, achievements, and forest progress.",
        "Responding to change over following a plan: when problems appeared, such as CI lockfile mismatch, TypeScript configuration issues, hardcoded secrets, Docker packaging failures, or Expo dependency drift, I adapted the plan and fixed the highest-impact blocker first.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Agile Rituals And Evidence", 1)
    add_table(
        doc,
        ["Ritual or practice", "My individual-project version", "Concrete evidence"],
        [
            (
                "Sprint planning",
                "At the start of each work phase I selected a small group of goals, such as authentication, CI quality, security hardening, or Docker setup.",
                "LEARNING_TRACKER.txt timeline and milestones",
            ),
            (
                "Daily/regular progress check",
                "I used the learning tracker and Git commits to keep track of what was done, what was blocked, and what needed attention next.",
                "LEARNING_TRACKER.txt and git commit history",
            ),
            (
                "Sprint review",
                "I verified increments through running code, tests, API checks, and documentation updates.",
                "Backend tests, Postman checks, README.md, CI workflow",
            ),
            (
                "Retrospective",
                "I wrote down solved issues and lessons, for example npm ci lock mismatch, Docker build failures, stale diagnostics, and Expo dependency mismatch.",
                "LEARNING_TRACKER.txt troubleshooting view",
            ),
            (
                "Definition of done",
                "A feature was done when it compiled, could be tested, was connected to the intended layer, and relevant setup/docs were updated.",
                "CI workflow, tests, README, .env.example",
            ),
        ],
        [1.45, 3.35, 1.70],
    )

    add_heading(doc, "Backlog Examples", 1)
    add_table(
        doc,
        ["Backlog item", "User or project value", "Status/evidence"],
        [
            (
                "Register and log in",
                "Users can create an account and access their own LifeForest data.",
                "User/Auth backend packages and frontend login/register screens",
            ),
            (
                "Create routines and tasks",
                "Users can structure their day and manage repeatable productivity work.",
                "Routine and Task backend/frontend files",
            ),
            (
                "Start focus sessions",
                "Users can run focused work sessions connected to their progress.",
                "FocusSession backend and frontend app flow",
            ),
            (
                "Reflect after focus work",
                "Users can record distraction and focus information for learning.",
                "Reflection backend and frontend flow",
            ),
            (
                "Show analytics and achievements",
                "Users receive feedback about their productivity and progress.",
                "Analytics and Achievements services/controllers/screens",
            ),
            (
                "Grow a forest",
                "The app visualises progress with trees to make the routine/focus loop motivating.",
                "Tree backend and forest frontend screens/components",
            ),
            (
                "Improve delivery quality",
                "The project is easier to test, run, deploy, and review.",
                "CI, Docker, README, security checks, portfolio evidence docs",
            ),
        ],
        [1.65, 3.00, 1.85],
    )

    add_heading(doc, "Adaptation Examples", 1)
    add_body(
        doc,
        "Agile development is valuable when the first plan turns out to be incomplete. The project shows several examples where I adjusted the work based on feedback from tools, runtime behaviour, or quality checks."
    )
    for item in [
        "When hardcoded secrets and configuration risks appeared, I moved database and JWT values into environment-driven configuration and added typed JwtProperties.",
        "When frontend API URLs were too rigid for phones, emulators, and Docker, I replaced hardcoded URLs with EXPO_PUBLIC_API_URL.",
        "When CI failed because of lockfile and runtime mismatches, I synced package-lock.json, updated GitHub Actions versions, and aligned Node/Expo dependencies.",
        "When Docker packaging failed because of duplicate Spring Boot metadata, I adjusted the build configuration and verified container startup.",
        "When the project needed stronger proof of quality, I added testing, security, and performance evidence documents in the docs folder.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Reflection", 1)
    add_body(
        doc,
        "Using Agile helped me keep the project manageable. Instead of trying to design the complete system perfectly at the start, I built a foundation, checked it, learned from problems, and then extended the product. This made the project more flexible and reduced the risk of discovering major issues only at the end."
    )
    add_body(
        doc,
        "The strongest evidence is the combination of the learning tracker, Git history, implemented feature packages, CI workflow, tests, Docker setup, and documentation. A future improvement would be to keep a dedicated agile board screenshot or export from a professional tool such as Jira, Azure DevOps, or GitHub Projects so that backlog movement and sprint planning are visible in a more standard industry format."
    )

    add_heading(doc, "Source Material Used", 1)
    for item in [
        "Agile Manifesto values and principles: https://agilemanifesto.org/",
        "Scrum Guide: https://scrumguides.org/",
        "Fontys Agile software development workshop criteria from the provided assignment text.",
        "LifeForest repository evidence: README.md, LEARNING_TRACKER.txt, .github/workflows/ci.yml, backend, frontend, docs, and Git history.",
    ]:
        add_bullet(doc, item)

    add_footer(doc.sections[0])
    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build_docx())
