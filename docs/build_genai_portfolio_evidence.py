from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "lifeforest-genai-evidence.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(25, 25, 25)
MUTED = RGBColor(90, 90, 90)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
BORDER = "C9D3DF"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=130, bottom=90, end=130):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    for index, width in enumerate(widths):
        table.columns[index].width = Inches(width)
        for cell in table.columns[index].cells:
            cell.width = Inches(width)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_cell_text(cell, text, bold=False, color=None, size=9.0):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.10
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color or INK


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    paragraph.add_run(text)
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.style = "Normal"
    paragraph.add_run(text)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.167
    paragraph.add_run(text)
    return paragraph


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.167
    paragraph.add_run(text)
    return paragraph


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    set_table_borders(table, color="D8E0EA", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
    title_para = cell.paragraphs[0]
    title_para.paragraph_format.space_after = Pt(3)
    title_run = title_para.add_run(title)
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(10.5)
    title_run.font.bold = True
    title_run.font.color.rgb = DARK_BLUE
    body_para = cell.add_paragraph()
    body_para.paragraph_format.space_after = Pt(0)
    body_run = body_para.add_run(text)
    body_run.font.name = "Calibri"
    body_run.font.size = Pt(10)
    body_run.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    mark_header_row(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        set_cell_text(cell, header, bold=True, color=DARK_BLUE, size=9.2)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row_data in rows:
        row = table.add_row()
        for index, text in enumerate(row_data):
            cell = row.cells[index]
            set_cell_shading(cell, WHITE)
            set_cell_margins(cell)
            set_cell_text(cell, text, bold=(index == 0), size=8.85)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    for list_style_name in ["List Bullet", "List Number"]:
        style = styles[list_style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.167

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("LifeForest - GenAI Portfolio Evidence")
    footer_run.font.name = "Calibri"
    footer_run.font.size = Pt(8.5)
    footer_run.font.color.rgb = MUTED


def build_docx():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    title_run = title.add_run("Generative AI in Software Development")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = BLUE

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle.add_run(
        "Portfolio evidence: how I explored and applied GenAI in the LifeForest project"
    )
    subtitle_run.font.name = "Calibri"
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = MUTED

    add_callout(
        doc,
        "Learning outcome claim",
        "I applied GenAI as an engineering support tool in LifeForest by using prompt-based assistance to analyse project evidence, structure documentation, identify improvement areas, and produce a proof-of-concept portfolio artifact. I stayed critical by checking the generated output against the repository, project files, and quality evidence instead of accepting AI output blindly.",
    )

    add_heading(doc, "Project Context", 1)
    add_body(
        doc,
        "LifeForest is a mobile-first productivity and focus platform with an Expo React Native frontend, a Spring Boot backend API, and a PostgreSQL database. The project includes authentication, routines, tasks, focus sessions, reflections, analytics, achievements, trees, CI, Docker setup, testing evidence, security evidence, and documentation."
    )
    add_body(
        doc,
        "The GenAI module asks students to explore how AI can support the software development lifecycle: analyse, design, realise, test, and deploy. I used LifeForest as the context for this exploration and treated GenAI as a tool that can speed up repetitive work, but still needs developer judgement, verification, and ethical use."
    )

    add_heading(doc, "Quick Scan Of GenAI Opportunities", 1)
    add_matrix(
        doc,
        ["SDLC phase", "Possible GenAI support", "LifeForest example"],
        [
            (
                "Analyse",
                "Summarise stakeholder/module criteria, extract user stories, find evidence gaps.",
                "Turn workshop text and repository facts into Portflow evidence sections.",
            ),
            (
                "Design",
                "Suggest architecture explanations, C4 diagram text, UI flow alternatives, or data-model questions.",
                "Use the existing Spring Boot, Expo, PostgreSQL, and C4 evidence as context for explanation.",
            ),
            (
                "Realise",
                "Generate boilerplate, DTO/service/controller patterns, helper scripts, or documentation builders.",
                "Create repeatable Python scripts that generate portfolio DOCX evidence from repository context.",
            ),
            (
                "Test",
                "Suggest unit-test cases, edge cases, mock data, and coverage gaps.",
                "Compare backend tests, frontend lint/typecheck, OWASP checks, and CI workflow against quality goals.",
            ),
            (
                "Deploy",
                "Draft deployment steps, environment-variable explanations, Docker troubleshooting, and README text.",
                "Explain Docker Compose, backend/frontend images, .env.example, and VM deployment instructions.",
            ),
        ],
        [1.2, 2.65, 2.65],
    )

    add_heading(doc, "Selected Tool And Proof Of Concept", 1)
    add_body(
        doc,
        "For the proof of concept I selected a large language model coding assistant, Codex, because it can read project files, inspect source code, run local commands, generate scripts, and create documentation artifacts. This matches the GenAI category of LLMs: tools that interpret, transform, and generate text or code based on prompts."
    )
    add_body(
        doc,
        "The proof of concept was to automate part of the portfolio documentation workflow. Instead of manually rewriting every workshop criterion into a separate evidence document, I used GenAI to transform the criteria into a project-specific DOCX document that explains how the concept was applied in LifeForest."
    )
    add_heading(doc, "Proof Of Concept Steps", 2)
    for item in [
        "Input: provide the workshop text about GenAI and ask for a Word document suitable for Portflow.",
        "Context gathering: inspect README.md, LEARNING_TRACKER.txt, .github/workflows/ci.yml, docs, backend, frontend, and Git history.",
        "Generation: create a Python DOCX builder under docs so the output is repeatable and stored with the project evidence.",
        "Project grounding: connect the generated document to LifeForest features, CI, testing, security, Docker, and documentation artifacts.",
        "Verification: structurally inspect the DOCX, check table/paragraph counts, and verify that it opens as a Word document artifact.",
        "Reflection: include limitations, ethical rules, and how I checked whether the AI output was actually correct.",
    ]:
        add_number(doc, item)

    add_heading(doc, "Prompting And Iteration", 1)
    add_body(
        doc,
        "The prompt was intentionally practical: make a Word document from the provided criteria that can be uploaded to Portflow and demonstrates how I applied it in this project. The useful prompt ingredients were the task, target audience, output format, and project context. The result was improved by iterating: first reading the criteria, then reading repository evidence, then generating a structured document instead of a generic summary."
    )
    add_matrix(
        doc,
        ["Prompting practice", "How I used it", "Why it improved the result"],
        [
            (
                "Give context",
                "I supplied the workshop criteria and allowed the assistant to inspect LifeForest files.",
                "The document can refer to real project artifacts instead of vague examples.",
            ),
            (
                "Ask for a concrete output",
                "The target was a Word document for Portflow, not just an explanation in chat.",
                "The output became a usable portfolio artifact.",
            ),
            (
                "Iterate",
                "The assistant first gathered context, then created a builder script, then generated the DOCX.",
                "The final result is more accurate and reusable than a one-shot answer.",
            ),
            (
                "Verify",
                "Generated statements were checked against README, learning tracker, CI, docs, and source tree.",
                "This reduces the risk of hallucinated evidence.",
            ),
        ],
        [1.45, 2.55, 2.50],
    )

    add_heading(doc, "Critical Evaluation Of The AI Output", 1)
    add_body(
        doc,
        "I did not treat the GenAI output as automatically correct. I checked whether the generated claims were supported by files in the repository. For example, the project really contains a backend, frontend, CI workflow, Docker documentation, security checks, testing evidence, and portfolio documents in the docs folder."
    )
    for item in [
        "Accuracy check: compare generated statements with README.md, LEARNING_TRACKER.txt, source packages, CI workflow, and docs artifacts.",
        "Scope check: avoid claiming that AI implemented every feature; the document only claims that AI was used as support for documentation, analysis, and evidence generation.",
        "Quality check: store the builder script so the document can be regenerated or improved instead of being a one-off black box.",
        "Risk check: avoid putting secrets, personal data, or private user data into prompts or generated output.",
        "Human responsibility: final decisions, project claims, and assessment evidence remain my responsibility as the developer.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Ethical And Legal Considerations", 1)
    add_matrix(
        doc,
        ["Consideration", "Risk", "How I handled it"],
        [
            (
                "Privacy",
                "Prompts can expose secrets, user data, or private project details.",
                "Use project files and public-style evidence only; do not paste secrets or real user data.",
            ),
            (
                "Academic integrity",
                "AI could produce work that hides the student's real understanding.",
                "Use AI as support and explain the process, verification, and personal reflection.",
            ),
            (
                "Hallucination",
                "The model may invent facts, tools, or evidence that are not in the repository.",
                "Check claims against concrete files before uploading to Portflow.",
            ),
            (
                "Security",
                "Generated code or advice can introduce vulnerabilities.",
                "Use CI, tests, OWASP checks, code review, and developer judgement before accepting output.",
            ),
            (
                "Copyright/licensing",
                "Generated content may be influenced by training data or external sources.",
                "Use short summaries, cite sources where relevant, and keep final project claims grounded in my own work.",
            ),
        ],
        [1.35, 2.35, 2.80],
    )

    add_heading(doc, "Impact On My Development Process", 1)
    add_body(
        doc,
        "The biggest benefit was speed in organising and explaining project evidence. GenAI helped convert broad criteria into a clear artifact structure, suggested which repository evidence mattered, and reduced the time needed to produce consistent documentation. This gave me more time to focus on checking whether the evidence was correct and complete."
    )
    add_body(
        doc,
        "The limitation is that GenAI does not replace understanding. It can write a convincing paragraph even when the underlying claim is not proven. That is why the proof of concept includes verification steps and why the generated document refers to concrete LifeForest artifacts instead of only describing GenAI in general terms."
    )

    add_heading(doc, "Evidence From LifeForest", 1)
    add_matrix(
        doc,
        ["Evidence item", "What it demonstrates", "Location"],
        [
            (
                "Generated GenAI evidence document",
                "A practical proof of concept where GenAI helps automate documentation work.",
                "docs/lifeforest-genai-evidence.docx",
            ),
            (
                "Reusable builder script",
                "The output is reproducible and can be edited like code.",
                "docs/build_genai_portfolio_evidence.py",
            ),
            (
                "Project README",
                "Shows the real architecture, setup, Docker workflow, and deployment notes used as grounding context.",
                "README.md",
            ),
            (
                "Learning tracker",
                "Shows the project timeline, issues solved, milestones, and lessons learned.",
                "LEARNING_TRACKER.txt",
            ),
            (
                "CI workflow",
                "Shows automated build, test, lint, typecheck, OWASP, and audit checks used for verification.",
                ".github/workflows/ci.yml",
            ),
            (
                "Other evidence documents",
                "Shows that the same documentation automation pattern can support several portfolio outcomes.",
                "docs/lifeforest-agile-evidence.docx and other docs files",
            ),
        ],
        [1.55, 3.05, 1.90],
    )

    add_heading(doc, "Conclusion", 1)
    add_body(
        doc,
        "This proof of concept shows that GenAI can support the software development process when it is used deliberately. In LifeForest, the strongest use case was not replacing the developer, but supporting analysis, documentation, evidence structuring, and quality reflection. The tool helped create a more polished and complete Portflow artifact, while the developer still had to provide context, verify accuracy, and decide what belonged in the final evidence."
    )
    add_body(
        doc,
        "My main lesson is that GenAI is most valuable when combined with an engineering mindset: prompt clearly, iterate, inspect the output, test claims, and keep responsibility with the human developer. AI is the tool; I am still responsible for the craft."
    )

    add_heading(doc, "Source Material Used", 1)
    for item in [
        "Fontys GenAI workshop criteria from the provided assignment text.",
        "LifeForest repository evidence: README.md, LEARNING_TRACKER.txt, .github/workflows/ci.yml, backend, frontend, docs, and Git history.",
        "General prompting and GenAI concepts from the workshop: LLMs, SLMs, RAG, diffusion models, prompt iteration, critical evaluation, and SDLC support.",
    ]:
        add_bullet(doc, item)

    doc.core_properties.author = "LifeForest"
    doc.core_properties.title = "LifeForest Generative AI Evidence"
    doc.core_properties.subject = "Portfolio evidence for GenAI in the software development process"
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_docx()
    print(DOCX_PATH)
