from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "lifeforest-business-context-evidence.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(25, 25, 25)
MUTED = RGBColor(90, 90, 90)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
BORDER = "C9D3DF"
WHITE = RGBColor(255, 255, 255)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_inches):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    for index, width in enumerate(widths_inches):
        table.columns[index].width = Inches(width)
        for cell in table.columns[index].cells:
            cell.width = Inches(width)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_cell_text(cell, text, bold=False, color=None, size=9.1):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.10
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    paragraph.add_run(text)
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.style = "Normal"
    paragraph.add_run(text)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.167
    paragraph.add_run(text)
    return paragraph


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.167
    paragraph.add_run(text)
    return paragraph


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    set_table_borders(table, color="D8E0EA", size="4")
    mark_header_row(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
    title_para = cell.paragraphs[0]
    title_para.paragraph_format.space_after = Pt(3)
    title_run = title_para.add_run(title)
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(10.5)
    title_run.font.bold = True
    title_run.font.color.rgb = DARK_BLUE
    body_para = cell.add_paragraph()
    body_para.paragraph_format.space_after = Pt(0)
    body_para.paragraph_format.line_spacing = 1.10
    body_run = body_para.add_run(text)
    body_run.font.name = "Calibri"
    body_run.font.size = Pt(10)
    body_run.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    mark_header_row(table.rows[0])
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_shading(header_cells[i], LIGHT_BLUE)
        set_cell_margins(header_cells[i])
        set_cell_text(header_cells[i], header, bold=True, color=DARK_BLUE, size=9.2)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_text(cells[i], value, size=8.9)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    for list_style_name in ["List Bullet", "List Number"]:
        style = styles[list_style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("LifeForest - Business Context Evidence")
    footer_run.font.name = "Calibri"
    footer_run.font.size = Pt(8.5)
    footer_run.font.color.rgb = MUTED


def build_docx():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    title_run = title.add_run("Business Context")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = BLUE

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle.add_run("Portfolio evidence: how I applied business context analysis to LifeForest")
    subtitle_run.font.name = "Calibri"
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = MUTED

    add_callout(
        doc,
        "Learning outcome claim",
        "I applied business context analysis by identifying the users, roles, data, and process steps around productivity planning before LifeForest, then modelling how the application changes the process by supporting routine planning, task management, focus sessions, reflections, analytics, achievements, and forest progress.",
    )

    add_heading(doc, "Project Context", 1)
    add_body(
        doc,
        "LifeForest is a mobile-first productivity and focus application. The application combines a React Native/Expo frontend, a Spring Boot backend API, and a PostgreSQL database. The business context I analysed is the routine-based productivity process used by a person who wants to plan tasks, stay focused, reflect on focus sessions, and understand progress over time.",
    )
    add_body(
        doc,
        "The application is not only a timer. It connects planning, execution, reflection, and feedback into one process. This matters because productivity support becomes more useful when the user can move from intention to action and then review the result without switching between separate tools.",
    )

    add_heading(doc, "Users, Roles, and Stakeholders", 1)
    add_matrix(
        doc,
        ["Role", "Need in the process", "How LifeForest supports it"],
        [
            ["Primary user", "Plan routines and tasks, run focus sessions, and keep motivation visible.", "Uses the mobile app to register/login, create routines, add tasks, start focus sessions, reflect, and view progress."],
            ["Productivity coach or mentor", "Needs a structured way to discuss habits, focus quality, and improvement over time.", "Can use the user's routines, reflections, analytics, achievements, and forest progress as discussion evidence."],
            ["Developer / project owner", "Needs to ensure the process is reliable, secure, and maintainable.", "Implements backend services, API endpoints, validation, testing, and deployment documentation."],
            ["Assessment stakeholder", "Needs evidence that the software fits a real business/process context.", "Reviews the before/after process model, data flow, role changes, and improvement conclusions."],
        ],
        [1.35, 2.5, 2.65],
    )

    add_heading(doc, "Current Business Context Before LifeForest", 1)
    add_body(
        doc,
        "Before introducing LifeForest, the productivity process is fragmented. A user may plan routines in a notes app, keep tasks in a to-do list, use a separate timer for focus sessions, and record reflections manually or not at all. Progress is therefore difficult to evaluate because the data is spread across different places or only exists in memory.",
    )
    add_heading(doc, "Existing Process Model", 2)
    for item in [
        "User decides what routine or task should be worked on.",
        "User manually writes the routine or task in a note, planner, or separate task app.",
        "User starts a timer in another tool or tracks time mentally.",
        "User completes or abandons the focus period without structured feedback.",
        "User may write a short reflection manually, but this is optional and disconnected from the task.",
        "User manually checks progress by looking through notes, task lists, or memory.",
    ]:
        add_number(doc, item)
    add_body(
        doc,
        "Information used or produced in this old process includes task names, routine descriptions, planned focus time, actual focus time, completion status, and reflection notes. The issue is that these information objects are not connected, which creates weak feedback and makes long-term improvement hard to measure.",
    )

    add_heading(doc, "Problems and Improvement Opportunities", 1)
    for item in [
        "Bottleneck: the user has to switch between multiple tools before, during, and after a focus session.",
        "Quality issue: reflection and progress data can be forgotten or recorded inconsistently.",
        "Motivation issue: progress is not visible in a rewarding way, so the user gets little immediate feedback.",
        "Data issue: tasks, sessions, reflections, analytics, achievements, and progress are not stored as one connected dataset.",
        "Process opportunity: automate session tracking and progress calculation while still keeping the user in control of planning and reflection.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "New Business Context With LifeForest", 1)
    add_body(
        doc,
        "After introducing LifeForest, the process becomes a supported digital workflow. The user still decides what to work on, but the application handles storage, linking, progress updates, analytics, achievements, and forest visualisation. This changes the process from manual tracking to guided self-management.",
    )
    add_heading(doc, "New Process Model", 2)
    for item in [
        "User registers or logs in, creating a personal account and protected identity in the system.",
        "User creates routines and tasks in the mobile app.",
        "User selects a task or focus goal and starts a focus session.",
        "LifeForest tracks the focus session and links it to the user, task, and progress data.",
        "After the session, the user can add a reflection about focus quality and experience.",
        "The backend updates related data such as trees, analytics, achievements, and habit/focus progress.",
        "User views analytics, achievements, and forest progress to understand improvement over time.",
    ]:
        add_number(doc, item)

    add_heading(doc, "Activities Automated or Supported by the Application", 1)
    add_matrix(
        doc,
        ["Activity", "Before LifeForest", "With LifeForest"],
        [
            ["Account and identity", "User identity may not exist, or data is stored locally in separate tools.", "Authentication and user-specific backend data connect actions to the correct user."],
            ["Routine planning", "Routine details are written manually in notes or task apps.", "Routines are created, updated, retrieved, and deleted through the app and API."],
            ["Task management", "Tasks are maintained separately from routines and focus time.", "Tasks belong to routines and can be managed as part of the same workflow."],
            ["Focus tracking", "Timer use is separate from task and progress data.", "Focus sessions are started through the app and stored in the backend."],
            ["Reflection", "Reflection is optional and disconnected.", "Reflections can be submitted after sessions and stored with user context."],
            ["Progress feedback", "Progress must be interpreted manually.", "Analytics, achievements, trees, and forest progress provide automatic feedback."],
        ],
        [1.45, 2.45, 2.6],
    )

    add_heading(doc, "Data and Documents in the Process", 1)
    add_matrix(
        doc,
        ["Data object", "Used or produced by", "Business value"],
        [
            ["User account and JWT token", "Login/register flow and secured API communication.", "Identifies the user and separates personal productivity data."],
            ["Routine", "Routine screens and routine backend service.", "Represents repeated behaviour the user wants to build."],
            ["Task", "Task screens and task backend service.", "Breaks routines into concrete actions that can be planned and completed."],
            ["Focus session", "Focus session screen and backend focus-session service.", "Records work periods and links effort to progress."],
            ["Reflection", "Reflection screen and backend reflection service.", "Captures qualitative feedback after focus work."],
            ["Analytics and achievements", "Analytics and achievement services.", "Turn stored activity into understandable progress evidence."],
            ["Tree / forest progress", "Tree service and forest visual components.", "Gives the user a visible motivational representation of focus consistency."],
        ],
        [1.55, 2.45, 2.5],
    )

    add_heading(doc, "Evidence From My Implementation", 1)
    add_body(
        doc,
        "The business context analysis is reflected directly in the project structure. The frontend contains screens for registration, login, routine creation/editing, task creation/editing, focus sessions, reflection, and analytics. The backend contains services and controllers for users, authentication, routines, tasks, focus sessions, reflections, analytics, achievements, habits, and trees.",
    )
    add_body(
        doc,
        "This shows that I translated the business process into software responsibilities. Planning is represented by routines and tasks, execution by focus sessions, evaluation by reflections and analytics, and motivation by achievements and tree/forest progress.",
    )

    add_heading(doc, "Conclusion", 1)
    add_body(
        doc,
        "By analysing the business context, I identified that the main value of LifeForest is not only storing tasks, but improving the whole productivity workflow. The application reduces tool switching, connects planning with focus execution, makes reflection more structured, and turns raw activity data into progress feedback. The role of the user changes from manually coordinating several disconnected tools to using one supported process where the application automates tracking and feedback while the user remains responsible for choosing goals and reflecting honestly.",
    )
    add_body(
        doc,
        "The biggest improvement is process continuity: data that was previously scattered is now connected through the user, routines, tasks, focus sessions, reflections, analytics, achievements, and forest progress. This makes LifeForest better aligned with the real business need: helping a person or coaching context understand and improve focus behaviour over time.",
    )

    doc.core_properties.author = "LifeForest"
    doc.core_properties.title = "LifeForest Business Context Evidence"
    doc.core_properties.subject = "Portfolio evidence for business context analysis"
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_docx()
    print(DOCX_PATH)
